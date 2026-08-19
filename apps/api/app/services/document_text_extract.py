"""从上传文件提取纯文本（知识库 / 通用）。"""
from __future__ import annotations

import html as html_module
import io
from pathlib import Path


_MIN_LEN = 10


def extract_text_from_bytes(filename: str, data: bytes) -> str:
    """按扩展名提取文本；失败或过短时抛 ValueError。"""
    ext = Path(filename or "").suffix.lower()
    if ext in (".txt", ".md", ".markdown", ".csv", ".json", ".log", ""):
        text = _decode_text(data)
    elif ext == ".pdf":
        text = _extract_pdf(data)
    elif ext == ".docx":
        text = _extract_docx(data)
    elif ext == ".doc":
        raise ValueError("暂不支持 .doc，请另存为 .docx 或导出为 PDF/TXT")
    else:
        raise ValueError(f"不支持的文件类型 {ext or '(无扩展名)'}，允许: .txt .md .pdf .docx")

    text = (text or "").strip()
    if len(text) < _MIN_LEN:
        raise ValueError("文件内容过短或无法解析为文本")
    return text


def _decode_text(data: bytes) -> str:
    for enc in ("utf-8-sig", "gb18030", "latin-1"):
        try:
            text = data.decode(enc)
            if "\x00" in text[:2000]:
                continue
            return text
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise ValueError("服务器未安装 pypdf，无法解析 PDF") from e
    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(parts)


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as e:
        raise ValueError("服务器未安装 python-docx，无法解析 DOCX") from e
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts)


def docx_to_preview_html(data: bytes) -> str:
    """将 docx 转为简单 HTML，供管理端/商家端在线预览。"""
    try:
        from docx import Document
    except ImportError as e:
        raise ValueError("服务器未安装 python-docx，无法预览 DOCX") from e
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        text = html_module.escape(p.text or "")
        if not text.strip():
            continue
        style = str(getattr(p.style, "name", "") or "")
        if "Heading" in style or style.startswith("标题"):
            level = 2
            if "1" in style:
                level = 1
            elif "3" in style:
                level = 3
            parts.append(f"<h{level}>{text}</h{level}>")
        else:
            parts.append(f"<p>{text}</p>")
    for table in doc.tables:
        parts.append('<table class="doc-table">')
        for row in table.rows:
            parts.append("<tr>")
            for cell in row.cells:
                parts.append(f"<td>{html_module.escape(cell.text or '')}</td>")
            parts.append("</tr>")
        parts.append("</table>")
    body = "".join(parts) if parts else "<p>（文档无文本内容）</p>"
    return (
        "<!DOCTYPE html><html><head><meta charset='utf-8'>"
        "<style>"
        "body{font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,sans-serif;"
        "padding:20px 24px;line-height:1.7;color:#1e293b;max-width:920px;margin:0 auto}"
        "h1,h2,h3{margin:1em 0 .5em}p{margin:.5em 0}"
        ".doc-table{border-collapse:collapse;margin:12px 0;width:100%}"
        ".doc-table td{border:1px solid #e2e8f0;padding:6px 10px;vertical-align:top}"
        "</style></head><body>"
        f"{body}</body></html>"
    )
